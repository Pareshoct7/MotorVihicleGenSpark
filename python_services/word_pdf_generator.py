#!/usr/bin/env python3
"""
Word Document and PDF Generator Service for Flutter Motor Vehicle Inspection App

This service:
1. Reads the Word template
2. Fills in inspection data
3. Checks/unchecks checkboxes based on inspection data
4. Saves the filled Word document
5. Converts Word document to PDF for sharing

Usage:
    python word_pdf_generator.py --inspection-json inspection_data.json --output output.pdf
"""

import sys
import json
import argparse
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import RGBColor, Pt
import subprocess
import os


class InspectionWordPDFGenerator:
    """Generates Word documents and PDFs from inspection data"""
    
    # Checkbox mapping for the Word document
    CHECKBOX_UNCHECKED = '□'
    CHECKBOX_CHECKED = '☑'
    
    def __init__(self, template_path):
        """Initialize with Word template path"""
        self.template_path = Path(template_path)
        if not self.template_path.exists():
            raise FileNotFoundError(f"Template not found: {template_path}")
    
    def generate_word_document(self, inspection_data, output_path):
        """
        Generate a filled Word document from inspection data
        
        Args:
            inspection_data (dict): Inspection data from Flutter app
            output_path (str): Path to save the filled Word document
        
        Returns:
            str: Path to the generated Word document
        """
        # Load the template
        doc = Document(self.template_path)
        
        # Fill in inspection details (Table 0)
        self._fill_inspection_details(doc, inspection_data)
        
        # Fill in checklist items (Table 1)
        self._fill_checklist_items(doc, inspection_data)
        
        # Fill in spare keys checkbox (Table 2)
        self._fill_spare_keys(doc, inspection_data)
        
        # Fill in signatures (Table 3)
        self._fill_signatures(doc, inspection_data)
        
        # Save the document
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        
        print(f"✅ Word document generated: {output_path}")
        return str(output_path)
    
    def _fill_inspection_details(self, doc, data):
        """Fill in the inspection details table (Table 0)"""
        table = doc.tables[0]
        
        # Row 0: Vehicle Registration No (Cell 1) and Store (Cell 3)
        if len(table.rows[0].cells) > 1:
            table.rows[0].cells[1].text = data.get('vehicleRegistrationNo', '')
        if len(table.rows[0].cells) > 3:
            # Show store name with store number if available
            store_name = data.get('storeName', '')
            store_number = data.get('storeNumber', '')
            if store_number:
                table.rows[0].cells[3].text = f"{store_name} ({store_number})"
            else:
                table.rows[0].cells[3].text = store_name
        
        # Row 1: Odometer Reading (Cell 1) and Date (Cell 3)
        if len(table.rows[1].cells) > 1:
            table.rows[1].cells[1].text = data.get('odometerReading', '')
        
        if len(table.rows[1].cells) > 3:
            # Format date
            inspection_date = data.get('inspectionDate', '')
            if inspection_date:
                try:
                    # Parse ISO date string
                    date_obj = datetime.fromisoformat(inspection_date.replace('Z', '+00:00'))
                    formatted_date = date_obj.strftime('%d/%m/%Y')
                    table.rows[1].cells[3].text = formatted_date
                except:
                    table.rows[1].cells[3].text = inspection_date
        
        # Row 2: Employee Name (Cell 1)
        if len(table.rows) > 2 and len(table.rows[2].cells) > 1:
            table.rows[2].cells[1].text = data.get('employeeName', '')
    
    def _fill_checklist_items(self, doc, data):
        """Fill in the checklist items table (Table 1) - EXACT template structure"""
        table = doc.tables[1]
        
        # Mapping based on EXACT template analysis
        # Format: (row_idx, col_idx, field_name)
        checklist_mapping = [
            # Row 1: Tyres (tread depth) [1,0] | Both tail lights [1,2]
            (1, 0, 'tyresTreadDepth'),
            (1, 2, 'tailLights'),
            
            # Row 2: Wheel nuts [2,0] | Headlights (low beam) [2,2]
            (2, 0, 'wheelNuts'),
            (2, 2, 'headlightsLowBeam'),
            
            # Row 3: (Outside header) | Headlights (high beam) [3,2]
            (3, 2, 'headlightsHighBeam'),
            
            # Row 4: Cleanliness [4,0] | Reverse lights [4,2]
            (4, 0, 'cleanliness'),
            (4, 2, 'reverseLights'),
            
            # Row 5: Body damage [5,0] | Brake lights [5,2]
            (5, 0, 'bodyDamage'),
            (5, 2, 'brakeLights'),
            
            # Row 6: Mirrors & Windows [6,0] | (Cab header)
            (6, 0, 'mirrorsWindows'),
            
            # Row 7: Signage [7,0] | Windscreen & wipers [7,2]
            (7, 0, 'signage'),
            (7, 2, 'windscreenWipers'),
            
            # Row 8: (Mechanical header) | Horn [8,2]
            (8, 2, 'horn'),
            
            # Row 9: Engine – oil & water [9,0] | Indicators [9,2]
            (9, 0, 'engineOilWater'),
            (9, 2, 'indicators'),
            
            # Row 10: Brakes [10,0] | Seat belts [10,2]
            (10, 0, 'brakes'),
            (10, 2, 'seatBelts'),
            
            # Row 11: Transmission [11,0] | Cleanliness (cab) [11,2]
            (11, 0, 'transmission'),
            (11, 2, 'cabCleanliness'),
            
            # Row 12: Checkboxes in [12,0], [12,1], [12,2] | Service log book [12,2]
            # Note: Row 12 has unusual structure with multiple checkboxes
            # The service log book checkbox is at [12,2]
            (12, 2, 'serviceLogBook'),
        ]
        
        for row_idx, col_idx, field_name in checklist_mapping:
            if row_idx < len(table.rows) and col_idx < len(table.rows[row_idx].cells):
                cell = table.rows[row_idx].cells[col_idx]
                checked = data.get(field_name, False)
                cell.text = self.CHECKBOX_CHECKED if checked else self.CHECKBOX_UNCHECKED
    
    def _fill_spare_keys(self, doc, data):
        """Fill in spare keys checkbox (Table 2)"""
        if len(doc.tables) > 2:
            table = doc.tables[2]
            if len(table.rows) > 0 and len(table.rows[0].cells) > 0:
                checked = data.get('spareKeys', False)
                table.rows[0].cells[0].text = self.CHECKBOX_CHECKED if checked else self.CHECKBOX_UNCHECKED
    
    def _fill_signatures(self, doc, data):
        """Fill in signature section (Table 3)"""
        if len(doc.tables) > 3:
            table = doc.tables[3]
            
            # Row 1: Employee signature (Cell 0)
            if len(table.rows) > 1 and len(table.rows[1].cells) > 0:
                if data.get('signature'):
                    table.rows[1].cells[0].text = data['signature']
            
            # Row 2: Dates (Cell 0 and Cell 2)
            if len(table.rows) > 2:
                inspection_date = data.get('inspectionDate', '')
                if inspection_date:
                    try:
                        date_obj = datetime.fromisoformat(inspection_date.replace('Z', '+00:00'))
                        formatted_date = date_obj.strftime('%d/%m/%Y')
                        
                        if len(table.rows[2].cells) > 0:
                            table.rows[2].cells[0].text = formatted_date
                        if len(table.rows[2].cells) > 2:
                            table.rows[2].cells[2].text = formatted_date
                    except:
                        pass
    
    def convert_to_pdf(self, word_path, pdf_path=None):
        """
        Convert Word document to PDF
        
        Args:
            word_path (str): Path to the Word document
            pdf_path (str): Path to save the PDF (optional, defaults to same name with .pdf)
        
        Returns:
            str: Path to the generated PDF
        """
        word_path = Path(word_path)
        
        if not word_path.exists():
            raise FileNotFoundError(f"Word document not found: {word_path}")
        
        # If no PDF path specified, use same name with .pdf extension
        if pdf_path is None:
            pdf_path = word_path.with_suffix('.pdf')
        else:
            pdf_path = Path(pdf_path)
        
        # Convert Word to PDF using LibreOffice
        try:
            # Ensure output directory exists
            pdf_path.parent.mkdir(parents=True, exist_ok=True)
            
            # Determine LibreOffice command (macOS uses 'soffice', Linux may use 'libreoffice')
            libreoffice_cmd = 'soffice'  # Default to macOS/standard command
            
            # Check if soffice is available, otherwise try libreoffice
            try:
                subprocess.run(['which', 'soffice'], capture_output=True, check=True)
            except subprocess.CalledProcessError:
                try:
                    subprocess.run(['which', 'libreoffice'], capture_output=True, check=True)
                    libreoffice_cmd = 'libreoffice'
                except subprocess.CalledProcessError:
                    raise FileNotFoundError("LibreOffice not found. Install with: brew install libreoffice (macOS) or apt-get install libreoffice (Linux)")
            
            # Use LibreOffice to convert
            result = subprocess.run([
                libreoffice_cmd,
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', str(pdf_path.parent),
                str(word_path)
            ], capture_output=True, text=True, timeout=30)
            
            if result.returncode == 0:
                # LibreOffice creates PDF with same name in output directory
                generated_pdf = pdf_path.parent / f"{word_path.stem}.pdf"
                
                # Rename if needed
                if generated_pdf != pdf_path:
                    generated_pdf.rename(pdf_path)
                
                print(f"✅ PDF generated: {pdf_path}")
                return str(pdf_path)
            else:
                print(f"❌ LibreOffice conversion failed: {result.stderr}")
                raise Exception(f"PDF conversion failed: {result.stderr}")
                
        except FileNotFoundError as e:
            print(f"❌ {str(e)}")
            raise
        except subprocess.TimeoutExpired:
            print("❌ PDF conversion timed out")
            raise
        except Exception as e:
            print(f"❌ Error converting to PDF: {e}")
            raise
    
    def generate_word_and_pdf(self, inspection_data, output_base_path):
        """
        Generate both Word document and PDF
        
        Args:
            inspection_data (dict): Inspection data from Flutter app
            output_base_path (str): Base path for output files (without extension)
        
        Returns:
            tuple: (word_path, pdf_path)
        """
        output_base = Path(output_base_path)
        
        # Generate Word document
        word_path = output_base.with_suffix('.docx')
        self.generate_word_document(inspection_data, word_path)
        
        # Convert to PDF
        pdf_path = output_base.with_suffix('.pdf')
        self.convert_to_pdf(word_path, pdf_path)
        
        return str(word_path), str(pdf_path)


def main():
    """Command-line interface for the generator"""
    parser = argparse.ArgumentParser(
        description='Generate Word documents and PDFs from inspection data'
    )
    parser.add_argument(
        '--template',
        required=True,
        help='Path to Word template file (.docx)'
    )
    parser.add_argument(
        '--inspection-json',
        required=True,
        help='Path to JSON file containing inspection data'
    )
    parser.add_argument(
        '--output',
        required=True,
        help='Output file path (without extension, will generate .docx and .pdf)'
    )
    parser.add_argument(
        '--word-only',
        action='store_true',
        help='Generate only Word document, skip PDF conversion'
    )
    
    args = parser.parse_args()
    
    # Load inspection data
    with open(args.inspection_json, 'r') as f:
        inspection_data = json.load(f)
    
    # Initialize generator
    generator = InspectionWordPDFGenerator(args.template)
    
    # Generate documents
    if args.word_only:
        word_path = generator.generate_word_document(
            inspection_data,
            args.output + '.docx'
        )
        print(f"\n✅ Success! Generated Word document: {word_path}")
    else:
        word_path, pdf_path = generator.generate_word_and_pdf(
            inspection_data,
            args.output
        )
        print(f"\n✅ Success!")
        print(f"   Word document: {word_path}")
        print(f"   PDF document: {pdf_path}")


if __name__ == '__main__':
    main()
